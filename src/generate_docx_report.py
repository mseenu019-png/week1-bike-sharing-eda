"""
Academic DOCX Report Generator for Bike Sharing Demand EDA Project
===================================================================
Phase 17: Produces `report/Week1_Bike_Sharing_EDA_Report.docx` with:
- Formal Academic Title Page & Executive Metadata
- Professional Heading Hierarchy & Formatting
- Formatted Data Tables with Shaded Headers
- Embedded High-Resolution Visualizations (Figures 1 to 8) with In-Depth Analytical Commentary
- 8+ Quantified Evidence-Based Insights
- Limitations, Future Scope, References, and Appendices
"""

import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import pandas as pd

def set_cell_background(cell, fill_hex):
    """Sets background shading color for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets padding for a table cell in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}>'
                      f'<w:top w:w="{top}" w:type="dxa"/>'
                      f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
                      f'<w:left w:w="{left}" w:type="dxa"/>'
                      f'<w:right w:w="{right}" w:type="dxa"/>'
                      f'</w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    """Applies subtle borders to a table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'<w:tblBorders {nsdecls("w")}>'
                        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
                        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
                        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
                        f'<w:insideV w:val="none"/>'
                        f'<w:left w:val="none"/>'
                        f'<w:right w:val="none"/>'
                        f'</w:tblBorders>')
    tblPr.append(borders)

def format_paragraph(p, space_before=0, space_after=6, line_spacing=1.15):
    """Applies clean typographic spacing to a paragraph."""
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    format_paragraph(p, space_before=18, space_after=8)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(27, 54, 93) # Navy Blue
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    format_paragraph(p, space_before=14, space_after=6)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(43, 92, 143) # Slate Blue
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    format_paragraph(p, space_before=10, space_after=4)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.italic = True
    run.font.color.rgb = RGBColor(60, 60, 60)
    return p

def add_body_p(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    format_paragraph(p, space_before=0, space_after=6, line_spacing=1.15)
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = "Calibri"
        r_bold.font.size = Pt(11)
        r_bold.font.bold = True
        r_bold.font.color.rgb = RGBColor(30, 30, 30)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(45, 45, 45)
    return p

def add_bullet_p(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    format_paragraph(p, space_before=0, space_after=4, line_spacing=1.15)
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = "Calibri"
        r_bold.font.size = Pt(11)
        r_bold.font.bold = True
        r_bold.font.color.rgb = RGBColor(30, 30, 30)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(45, 45, 45)
    return p

def add_callout(doc, text, title="ACADEMIC NOTE"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F0F4F8")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="1D3557"/>'
                        f'<w:top w:val="none"/>'
                        f'<w:bottom w:val="none"/>'
                        f'<w:right w:val="none"/>'
                        f'</w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    format_paragraph(p, space_before=0, space_after=2)
    r_t = p.add_run(f"[{title}] ")
    r_t.font.name = "Calibri"
    r_t.font.bold = True
    r_t.font.size = Pt(10.5)
    r_t.font.color.rgb = RGBColor(29, 53, 87)
    
    r_body = p.add_run(text)
    r_body.font.name = "Calibri"
    r_body.font.size = Pt(10.5)
    r_body.font.color.rgb = RGBColor(50, 50, 50)
    
    p_after = doc.add_paragraph()
    format_paragraph(p_after, space_before=0, space_after=4)

def insert_figure(doc, img_path, fig_num, title, purpose, observation, interpretation):
    if not os.path.exists(img_path):
        print(f"[WARNING] Image missing: {img_path}")
        return
        
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(p_img, space_before=10, space_after=4)
    p_img.add_run().add_picture(img_path, width=Inches(5.8))
    
    # Caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(p_cap, space_before=2, space_after=8)
    r_cap = p_cap.add_run(f"Figure {fig_num}: {title}")
    r_cap.font.name = "Calibri"
    r_cap.font.size = Pt(10)
    r_cap.font.bold = True
    r_cap.font.color.rgb = RGBColor(27, 54, 93)
    
    # Structural breakdown
    add_bullet_p(doc, purpose, bold_prefix="• Purpose: ")
    add_bullet_p(doc, observation, bold_prefix="• Key Observation: ")
    add_bullet_p(doc, interpretation, bold_prefix="• Analytical Interpretation: ")
    
    p_sep = doc.add_paragraph()
    format_paragraph(p_sep, space_before=0, space_after=4)

def build_docx_report(base_dir: str = None):
    if base_dir is None:
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        
    report_dir = os.path.join(base_dir, "report")
    viz_dir = os.path.join(base_dir, "visualizations")
    output_docx_path = os.path.join(report_dir, "Week1_Bike_Sharing_EDA_Report.docx")
    
    os.makedirs(report_dir, exist_ok=True)
    doc = Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
        
    print("=" * 70)
    print("PHASE 17: GENERATING ACADEMIC DOCX REPORT")
    print(f"Target: {output_docx_path}")
    print("=" * 70)
    
    # -------------------------------------------------------------
    # COVER / TITLE PAGE
    # -------------------------------------------------------------
    p_space1 = doc.add_paragraph()
    format_paragraph(p_space1, space_before=36, space_after=0)
    
    p_main_title = doc.add_paragraph()
    p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(p_main_title, space_before=0, space_after=8)
    r_title = p_main_title.add_run("Data Acquisition, Cleaning and Exploratory Data Analysis of Bike Sharing Demand")
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(27, 54, 93) # Primary Navy
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(p_sub, space_before=0, space_after=24)
    r_sub = p_sub.add_run("A Python-Based Data Preparation and Exploratory Analysis Project")
    r_sub.font.name = "Calibri"
    r_sub.font.size = Pt(14)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(80, 80, 80)
    
    # Decorative rule
    p_rule = doc.add_paragraph()
    p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(p_rule, space_before=0, space_after=36)
    r_line = p_rule.add_run("―" * 40)
    r_line.font.color.rgb = RGBColor(180, 180, 180)
    
    p_meta_box = doc.add_paragraph()
    p_meta_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(p_meta_box, space_before=0, space_after=4)
    r_meta_h = p_meta_box.add_run("ACADEMIC PORTFOLIO REPORT — WEEK 1\n")
    r_meta_h.font.name = "Calibri"
    r_meta_h.font.size = Pt(12)
    r_meta_h.font.bold = True
    r_meta_h.font.color.rgb = RGBColor(29, 53, 87)
    
    # Metadata Table
    meta_table = doc.add_table(rows=7, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Student Name:", "[Student Name]"),
        ("Register / Student ID:", "[Register Number]"),
        ("Department:", "Department of Computer Science & Data Analytics"),
        ("Institution:", "Faculty of Computing and Information Sciences"),
        ("Course / Track:", "Data Science Internship Program — Week 1 Portfolio"),
        ("Dataset:", "UCI Machine Learning Repository: Bike Sharing Dataset (hour.csv)"),
        ("Submission Date:", "August 24, 2026")
    ]
    
    for row_idx, (label, val) in enumerate(meta_data):
        cell_lbl = meta_table.cell(row_idx, 0)
        cell_val = meta_table.cell(row_idx, 1)
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(3.8)
        
        p_l = cell_lbl.paragraphs[0]
        format_paragraph(p_l, space_before=2, space_after=2)
        rl = p_l.add_run(label)
        rl.font.name = "Calibri"
        rl.font.bold = True
        rl.font.size = Pt(10.5)
        rl.font.color.rgb = RGBColor(40, 40, 40)
        
        p_v = cell_val.paragraphs[0]
        format_paragraph(p_v, space_before=2, space_after=2)
        rv = p_v.add_run(val)
        rv.font.name = "Calibri"
        rv.font.size = Pt(10.5)
        rv.font.color.rgb = RGBColor(60, 60, 60)
        
    set_table_borders(meta_table, color="D0D7DE")
    
    doc.add_page_break()
    
    # -------------------------------------------------------------
    # TABLE OF CONTENTS
    # -------------------------------------------------------------
    add_heading_1(doc, "TABLE OF CONTENTS")
    toc_items = [
        ("1. Introduction and Project Overview", "3"),
        ("2. Dataset Selection and Description", "4"),
        ("3. Data Acquisition and Ingestion Architecture", "5"),
        ("4. Initial Dataset Understanding and Schema Profiling", "6"),
        ("5. Comprehensive Data Quality Audit", "7"),
        ("   5.1 Missing Value Verification", "7"),
        ("   5.2 Duplicate Record Analysis", "7"),
        ("   5.3 Data Type Correctness", "8"),
        ("   5.4 Logical Range Validation and Domain Assertions", "8"),
        ("   5.5 Outlier Assessment and Domain Rationale", "9"),
        ("6. Data Cleaning and Preprocessing Pipeline", "10"),
        ("7. Feature Engineering for Exploratory Analysis", "11"),
        ("8. Exploratory Data Analysis & Statistical Synthesis", "12"),
        ("   8.1 Descriptive Statistics Summary", "12"),
        ("   8.2 Overall Demand Distribution Analysis", "13"),
        ("   8.3 Diurnal Commute and Hourly Dynamics", "13"),
        ("   8.4 Monthly Seasonality and System Growth", "14"),
        ("   8.5 Seasonal Climatological Patterns", "14"),
        ("   8.6 Environmental Drivers & Weather Severity", "15"),
        ("   8.7 Working-Day vs Non-Working-Day User Segmentation", "16"),
        ("   8.8 Multi-Variable Correlation Analysis", "17"),
        ("9. Publication-Grade Visualizations & Analytical Interpretations", "18"),
        ("10. Key Quantified Insights", "23"),
        ("11. Limitations of the Study", "24"),
        ("12. Future Scope and Predictive Opportunities", "25"),
        ("13. Conclusion", "26"),
        ("14. References", "27"),
        ("Appendix A: Complete Python Source Architecture", "28"),
        ("Appendix B: Data Quality Audit and Transformation Matrix", "29")
    ]
    
    toc_table = doc.add_table(rows=len(toc_items), cols=2)
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (section, page_str) in enumerate(toc_items):
        c0 = toc_table.cell(idx, 0)
        c1 = toc_table.cell(idx, 1)
        c0.width = Inches(5.5)
        c1.width = Inches(1.0)
        
        p0 = c0.paragraphs[0]
        format_paragraph(p0, space_before=1, space_after=1)
        r0 = p0.add_run(section)
        r0.font.name = "Calibri"
        r0.font.size = Pt(10.5)
        if not section.startswith("   "):
            r0.font.bold = True
            r0.font.color.rgb = RGBColor(27, 54, 93)
        else:
            r0.font.color.rgb = RGBColor(70, 70, 70)
            
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        format_paragraph(p1, space_before=1, space_after=1)
        r1 = p1.add_run(page_str)
        r1.font.name = "Calibri"
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = RGBColor(100, 100, 100)
        
    set_table_borders(toc_table, color="E5E5E5")
    
    doc.add_page_break()
    
    # -------------------------------------------------------------
    # SECTION 1: INTRODUCTION
    # -------------------------------------------------------------
    add_heading_1(doc, "1. INTRODUCTION")
    add_body_p(doc, "Data Science is the systematic, multi-disciplinary paradigm that combines domain expertise, statistical mathematics, algorithm development, and data engineering to extract robust empirical insights and actionable operational knowledge from raw, structured, and unstructured observations.")
    add_body_p(doc, "In the modern industrial and academic data science lifecycle, Data Preparation (Acquisition, Auditing, Cleaning, and Validation) and Exploratory Data Analysis (EDA) constitute the indispensable foundational phases. Empirical studies across industrial machine learning deployments indicate that data engineering and quality assessment represent upwards of 70% to 80% of total project effort. Jumping prematurely into complex machine learning models without rigorous data auditing invariably leads to flawed inferences, data leakage, model bias, and severe operational degradation ('garbage in, garbage out').")
    add_body_p(doc, "Exploratory Data Analysis, originally formalized by John W. Tukey, serves as the primary investigative apparatus for examining dataset distributions, discovering underlying geometric structures, identifying anomalies and outliers, verifying domain assumptions, testing operational hypotheses, and formulating feature engineering strategies.")
    add_body_p(doc, "Project Objectives: This Week 1 project establishes a professional, reproducible data preparation and exploratory analysis framework applied to the real-world UCI Bike Sharing Dataset. We execute a disciplined end-to-end pipeline encompassing programmatic ingestion, comprehensive data quality audits, boundary constraint validation, conservative preprocessing, feature derivation, multi-dimensional statistical synthesis, and publication-standard visualization.")
    
    add_callout(doc, "This project strictly simulates an authentic enterprise data preparation workflow. In compliance with data science ethics and academic integrity, no synthetic missing values or simulated anomalies were artificially introduced into the pristine UCI dataset; instead, exhaustive audits and logical assertions verify and document data integrity.", "METHODOLOGICAL PRINCIPLE")

    # -------------------------------------------------------------
    # SECTION 2: DATASET SELECTION
    # -------------------------------------------------------------
    add_heading_1(doc, "2. DATASET SELECTION AND DESCRIPTION")
    add_body_p(doc, "The project utilizes the official, publicly available Bike Sharing Dataset hosted by the University of California, Irvine (UCI) Machine Learning Repository. The dataset captures the complete operational telemetry of the Capital Bikeshare system in Washington, D.C., USA, spanning a continuous two-year window from January 1, 2011 to December 31, 2012.")
    
    add_bullet_p(doc, "UCI Machine Learning Repository: Bike Sharing Dataset", bold_prefix="• Official Name: ")
    add_bullet_p(doc, "https://archive.ics.uci.edu/dataset/275/bike+sharing+dataset", bold_prefix="• Official Repository URL: ")
    add_bullet_p(doc, "https://doi.org/10.24432/C5W894 (Digital Object Identifier)", bold_prefix="• Dataset DOI: ")
    add_bullet_p(doc, "Hadi Fanaee-T and Joao Gama (Laboratory of Artificial Intelligence and Decision Support, University of Porto)", bold_prefix="• Primary Curators: ")
    add_bullet_p(doc, "17,379 observation records across 17 raw attributes.", bold_prefix="• Selected Granularity: ")
    add_bullet_p(doc, "Total hourly rental demand (cnt), decomposed additively into casual recreational users (casual) and registered annual subscribers (registered).", bold_prefix="• Response Target: ")
    
    add_body_p(doc, "Rationale for Dataset Selection: Bike sharing systems are highly sensitive to environmental factors (temperature, relative humidity, precipitation, windspeed), calendar cycles (hour of the day, day of the week, working day vs weekend, month, and season), and distinct user behavioral profiles (commuters vs leisure tourists). This makes the dataset ideal for mastering EDA techniques.")

    # -------------------------------------------------------------
    # SECTION 3: DATA ACQUISITION
    # -------------------------------------------------------------
    add_heading_1(doc, "3. DATA ACQUISITION AND INGESTION ARCHITECTURE")
    add_body_p(doc, "The data acquisition architecture is implemented in 'src/data_acquisition.py'. It programmatically downloads the compressed official zip archive directly from the UCI repository endpoint (https://archive.ics.uci.edu/static/public/275/bike+sharing+dataset.zip), verifies archive contents (['hour.csv', 'day.csv', 'Readme.txt']), extracts 'hour.csv' into 'data/raw/hour.csv', and validates file integrity.")
    add_body_p(doc, "Raw Data Preservation Policy: In accordance with professional data engineering governance, raw source files are treated as immutable read-only artifacts. All subsequent cleaning, type casting, and feature engineering transformations are performed on explicit DataFrame copies and persisted to separate storage in 'data/processed/bike_sharing_cleaned.csv'.")

    # -------------------------------------------------------------
    # SECTION 4: DATASET UNDERSTANDING & SCHEMA PROFILING
    # -------------------------------------------------------------
    add_heading_1(doc, "4. INITIAL DATASET UNDERSTANDING AND SCHEMA PROFILING")
    add_body_p(doc, "The raw dataset consists of exactly 17,379 rows and 17 columns with a memory footprint of approximately 2.3 MB. The table below provides the complete data dictionary.")
    
    # Table of Schema
    schema_data = [
        ("instant", "int64", "Unique row sequential index", "1 to 17379"),
        ("dteday", "object", "Date of observation (YYYY-MM-DD)", "2011-01-01 to 2012-12-31"),
        ("season", "int64", "Season category (1:Spring, 2:Summer, 3:Fall, 4:Winter)", "1, 2, 3, 4"),
        ("yr", "int64", "Year indicator (0: 2011, 1: 2012)", "0, 1"),
        ("mnth", "int64", "Month of the year", "1 to 12"),
        ("hr", "int64", "Hour of the day", "0 to 23"),
        ("holiday", "int64", "Binary indicator if day is official holiday", "0 (No), 1 (Yes)"),
        ("weekday", "int64", "Day of the week (0:Sunday, 1:Monday ... 6:Saturday)", "0 to 6"),
        ("workingday", "int64", "1 if day is neither weekend nor holiday; else 0", "0 (Non-Work), 1 (Work)"),
        ("weathersit", "int64", "Weather condition severity category (1 to 4)", "1 (Clear) to 4 (Severe)"),
        ("temp", "float64", "Normalized temperature in Celsius (divided by 41)", "0.02 to 1.00"),
        ("atemp", "float64", "Normalized feeling temperature in Celsius (divided by 50)", "0.00 to 1.00"),
        ("hum", "float64", "Normalized relative humidity (divided by 100)", "0.00 to 1.00"),
        ("windspeed", "float64", "Normalized wind speed (divided by 67)", "0.00 to 0.85"),
        ("casual", "int64", "Count of casual / non-registered users", "0 to 367"),
        ("registered", "int64", "Count of registered subscription users", "0 to 886"),
        ("cnt", "int64", "Total bike rentals (cnt = casual + registered)", "1 to 977")
    ]
    
    tbl_schema = doc.add_table(rows=len(schema_data) + 1, cols=4)
    tbl_schema.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Feature Name", "Raw Data Type", "Semantic Description", "Observed Range"]
    
    for c_idx, h_text in enumerate(headers):
        cell = tbl_schema.cell(0, c_idx)
        set_cell_background(cell, "1D3557")
        p = cell.paragraphs[0]
        format_paragraph(p, space_before=3, space_after=3)
        r = p.add_run(h_text)
        r.font.name = "Calibri"
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    for r_idx, row in enumerate(schema_data):
        for c_idx, val in enumerate(row):
            cell = tbl_schema.cell(r_idx + 1, c_idx)
            set_cell_background(cell, "FFFFFF" if r_idx % 2 == 0 else "F8F9FA")
            p = cell.paragraphs[0]
            format_paragraph(p, space_before=2, space_after=2)
            r = p.add_run(val)
            r.font.name = "Calibri"
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(40, 40, 40)
            
    set_table_borders(tbl_schema, color="D0D7DE")
    
    doc.add_page_break()

    # -------------------------------------------------------------
    # SECTION 5: DATA QUALITY ASSESSMENT
    # -------------------------------------------------------------
    add_heading_1(doc, "5. COMPREHENSIVE DATA QUALITY AUDIT")
    add_body_p(doc, "A rigorous, structured data quality audit was conducted across 5 foundational dimensions. Every audit item follows the structured framework: Problem → Detection Method → Decision → Justification → Result.")
    
    add_heading_2(doc, "5.1 Missing Value Verification")
    add_bullet_p(doc, "Potential presence of null, NaN, or unrecorded telemetry entries in raw records.", bold_prefix="• Problem: ")
    add_bullet_p(doc, "Executed `raw_df.isnull().sum()` across all 17 features.", bold_prefix="• Detection Method: ")
    add_bullet_p(doc, "No imputation or deletion necessary.", bold_prefix="• Decision: ")
    add_bullet_p(doc, "Zero missing observations were detected (0 null values across all 17,379 rows). Imputing synthetic data would corrupt benchmark integrity.", bold_prefix="• Justification: ")
    add_bullet_p(doc, "Dataset confirmed 100.0% complete across all features.", bold_prefix="• Result: ")
    
    add_heading_2(doc, "5.2 Duplicate Record Analysis")
    add_bullet_p(doc, "Risk of duplicated observation records or non-unique identifier keys.", bold_prefix="• Problem: ")
    add_bullet_p(doc, "Evaluated `raw_df.duplicated().sum()` and primary key uniqueness `raw_df['instant'].nunique()`.", bold_prefix="• Detection Method: ")
    add_bullet_p(doc, "Retain all 17,379 records without deduplication filtering.", bold_prefix="• Decision: ")
    add_bullet_p(doc, "Zero exact duplicate rows detected; instant index strictly monotonic from 1 to 17,379.", bold_prefix="• Justification: ")
    add_bullet_p(doc, "Confirmed 100% uniqueness of records.", bold_prefix="• Result: ")
    
    add_heading_2(doc, "5.3 Incorrect Data Types")
    add_bullet_p(doc, "The date column `dteday` was stored as generic string object, preventing temporal indexing.", bold_prefix="• Problem: ")
    add_bullet_p(doc, "Inspected `raw_df.dtypes`.", bold_prefix="• Detection Method: ")
    add_bullet_p(doc, "Cast `dteday` into proper `datetime64[ns]` object.", bold_prefix="• Decision: ")
    add_bullet_p(doc, "Enables precise time-series operations, calendar filtering, and period aggregations.", bold_prefix="• Justification: ")
    add_bullet_p(doc, "Date successfully converted to datetime format.", bold_prefix="• Result: ")

    add_heading_2(doc, "5.4 Logical Range Validation & Domain Assertions")
    add_body_p(doc, "Eleven programmatic assertions were executed to enforce physical and logical domain boundaries:")
    add_bullet_p(doc, "assert (casual + registered == cnt).all() -> 0 discrepancies across 17,379 rows.", bold_prefix="• Additive Consistency: ")
    add_bullet_p(doc, "assert hr.between(0, 23).all() and mnth.between(1, 12).all() -> 100% valid.", bold_prefix="• Temporal Ranges: ")
    add_bullet_p(doc, "assert season.between(1, 4).all() and weathersit.between(1, 4).all() -> 100% valid.", bold_prefix="• Categorical Bounds: ")
    add_bullet_p(doc, "assert (cnt >= 0).all() and temp.between(0, 1).all() and hum.between(0, 1).all() -> 100% valid.", bold_prefix="• Physical Quantities: ")
    
    add_heading_2(doc, "5.5 Outlier Assessment & Domain Rationale")
    add_body_p(doc, "Using the standard Tukey Interquartile Range (IQR) fence ($[Q_1 - 1.5 \\times \\text{IQR}, Q_3 + 1.5 \\times \\text{IQR}]$), potential statistical outliers were identified:")
    add_bullet_p(doc, "505 observations (2.91%) exceed the upper statistical fence of 642.5 rentals/hr (maximum: 977).", bold_prefix="• Total Rentals (cnt): ")
    add_bullet_p(doc, "1,192 observations (6.86%) exceed upper fence of 114 rentals/hr.", bold_prefix="• Casual Users (casual): ")
    add_bullet_p(doc, "680 observations (3.91%) exceed upper fence of 499 rentals/hr.", bold_prefix="• Registered Users (registered): ")
    add_bullet_p(doc, "0 outliers for temperature; 22 (0.13%) for humidity near 0%; 342 (1.97%) for windspeed.", bold_prefix="• Environmental Features: ")
    
    add_callout(doc, "CRITICAL OUTLIER DECISION: Outlier records in rental demand (cnt > 642.5) occur systematically during peak rush hours (17:00-18:00) on clear weekdays and during high-volume summer holiday events. These represent legitimate real-world surge demand rather than telemetry errors. Deleting them would artificially truncate operational peaks and severely distort demand forecasting.", "ANALYTICAL GOVERNANCE")

    # -------------------------------------------------------------
    # SECTION 6 & 7: DATA CLEANING & FEATURE ENGINEERING
    # -------------------------------------------------------------
    add_heading_1(doc, "6. DATA CLEANING AND PREPROCESSING PIPELINE")
    add_body_p(doc, "Data cleaning was executed modularly in 'src/data_cleaning.py'. A separate DataFrame `clean_df` was generated while keeping the raw dataset untouched. Transformations included datetime casting, categorical mapping, and physical scale reconstruction.")

    add_heading_1(doc, "7. FEATURE ENGINEERING FOR EXPLORATORY ANALYSIS")
    add_body_p(doc, "To facilitate intuitive exploratory visualization and statistical grouping without ML encoding, ten descriptive features were derived:")
    
    feat_data = [
        ("season_name", "season (1..4)", "Categorical string ('Spring', 'Summer', 'Fall', 'Winter')", "Categorical grouping"),
        ("year", "yr (0, 1)", "Calendar year (2011, 2012)", "Annual growth tracking"),
        ("month_name", "mnth (1..12)", "Full month string ('January' .. 'December')", "Monthly seasonality"),
        ("day_of_week", "weekday (0..6)", "Day name string ('Sunday' .. 'Saturday')", "Weekly demand rhythm"),
        ("weather_description", "weathersit (1..4)", "Human-readable weather status ('Clear / Few Clouds' ...)", "Weather impact"),
        ("hour_group", "hr (0..23)", "Diurnal day-part ('Morning Rush (06-09)', 'Evening Rush (16-19)'...)", "Commuter period binning"),
        ("is_weekend", "weekday", "Binary indicator (1 if Saturday/Sunday, else 0)", "Leisure vs work split"),
        ("temp_celsius", "temp", "Un-normalized temperature in °C ($temp \\times 41$)", "Real-world physical scale"),
        ("atemp_celsius", "atemp", "Un-normalized feeling temperature in °C ($atemp \\times 50$)", "Human thermal comfort"),
        ("humidity_pct", "hum", "Un-normalized relative humidity percentage ($hum \\times 100$)", "Moisture measurement"),
        ("windspeed_kmh", "windspeed", "Un-normalized wind speed in km/h ($windspeed \\times 67$)", "Aerodynamic resistance")
    ]
    
    tbl_feat = doc.add_table(rows=len(feat_data) + 1, cols=4)
    tbl_feat.alignment = WD_TABLE_ALIGNMENT.CENTER
    feat_headers = ["Engineered Feature", "Source Variable", "Transformation Formula / Logic", "Analytical Purpose"]
    
    for c_idx, h_text in enumerate(feat_headers):
        cell = tbl_feat.cell(0, c_idx)
        set_cell_background(cell, "1D3557")
        p = cell.paragraphs[0]
        format_paragraph(p, space_before=3, space_after=3)
        r = p.add_run(h_text)
        r.font.name = "Calibri"
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    for r_idx, row in enumerate(feat_data):
        for c_idx, val in enumerate(row):
            cell = tbl_feat.cell(r_idx + 1, c_idx)
            set_cell_background(cell, "FFFFFF" if r_idx % 2 == 0 else "F8F9FA")
            p = cell.paragraphs[0]
            format_paragraph(p, space_before=2, space_after=2)
            r = p.add_run(val)
            r.font.name = "Calibri"
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(40, 40, 40)
            
    set_table_borders(tbl_feat, color="D0D7DE")
    
    doc.add_page_break()

    # -------------------------------------------------------------
    # SECTION 8: EXPLORATORY DATA ANALYSIS & SUMMARY STATISTICS
    # -------------------------------------------------------------
    add_heading_1(doc, "8. EXPLORATORY DATA ANALYSIS & STATISTICAL SYNTHESIS")
    add_body_p(doc, "Comprehensive parametric and non-parametric summary statistics were calculated across all continuous variables to profile location, spread, dispersion, and distribution shapes.")
    
    # Summary stats table
    df_stats = pd.read_csv(os.path.join(base_dir, "outputs", "summary_statistics.csv"))
    tbl_s = doc.add_table(rows=len(df_stats) + 1, cols=9)
    tbl_s.alignment = WD_TABLE_ALIGNMENT.CENTER
    s_headers = ["Variable", "Mean", "Std Dev", "Median", "IQR", "Min", "Max", "Skew", "Kurtosis"]
    
    for c_idx, h_text in enumerate(s_headers):
        cell = tbl_s.cell(0, c_idx)
        set_cell_background(cell, "1D3557")
        p = cell.paragraphs[0]
        format_paragraph(p, space_before=3, space_after=3)
        r = p.add_run(h_text)
        r.font.name = "Calibri"
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    for r_idx, row in df_stats.iterrows():
        vals = [
            row["Variable"], str(row["Mean"]), str(row["Std_Dev"]), str(row["Median"]),
            str(row["IQR"]), str(row["Min"]), str(row["Max"]), str(row["Skewness"]), str(row["Kurtosis"])
        ]
        for c_idx, val in enumerate(vals):
            cell = tbl_s.cell(r_idx + 1, c_idx)
            set_cell_background(cell, "FFFFFF" if r_idx % 2 == 0 else "F8F9FA")
            p = cell.paragraphs[0]
            format_paragraph(p, space_before=2, space_after=2)
            r = p.add_run(val)
            r.font.name = "Calibri"
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(40, 40, 40)
            
    set_table_borders(tbl_s, color="D0D7DE")
    
    add_heading_2(doc, "8.1 Analytical Synthesis of Statistical Distributions")
    add_bullet_p(doc, "Total hourly demand exhibits strong positive right-skewness (+1.19) with a mean of 189.46 rentals/hr and a median of 142.00 rentals/hr. The interquartile range spans 241.00 rentals/hr (Q1: 40.00 to Q3: 281.00).", bold_prefix="• Demand Skewness: ")
    add_bullet_p(doc, "Registered users show moderate skewness (+1.18, mean: 153.79), whereas casual demand is highly skewed (+2.50, mean: 35.67, median: 17.00), demonstrating that casual usage is concentrated in episodic leisure bursts.", bold_prefix="• User Sub-populations: ")
    add_bullet_p(doc, "Ambient temperature averages 20.38°C (std: 7.89°C, range: 0.82°C to 41.00°C) with nearly symmetric distribution (skewness: -0.01). Relative humidity averages 62.72% (skewness: -0.09).", bold_prefix="• Environmental Stability: ")

    doc.add_page_break()

    # -------------------------------------------------------------
    # SECTION 9: VISUALIZATIONS & DEEP ANALYTICAL INTERPRETATION
    # -------------------------------------------------------------
    add_heading_1(doc, "9. PUBLICATION-GRADE VISUALIZATIONS & ANALYTICAL INTERPRETATIONS")
    add_body_p(doc, "In accordance with professional standards, all eight visualizations are presented below with their dedicated figure caption, methodological purpose, key empirical observation, and analytical interpretation.")
    
    # Figure 1
    insert_figure(
        doc,
        os.path.join(viz_dir, "missing_values.png"),
        1,
        "Data Quality Audit — Feature Completeness",
        "Verify absence of missing, NaN, or null values across all 17 features in the dataset.",
        "All 17 features exhibit 100.0% data completeness (0 missing observations across 17,379 rows).",
        "The official UCI Bike Sharing Dataset is exceptionally clean in terms of completeness, eliminating the need for synthetic data imputation and ensuring unbiased statistical evaluation."
    )
    
    # Figure 2
    insert_figure(
        doc,
        os.path.join(viz_dir, "demand_distribution.png"),
        2,
        "Distribution of Hourly Bike Rental Demand",
        "Analyze the central tendency, dispersion, spread, and shape of total hourly bike rentals (cnt).",
        "The distribution exhibits strong right-skewness (+1.19) with a sharp peak in low-demand hours (overnight 00:00-05:00) and a long tail extending up to 977 rentals/hr. The mean (189.5) sits substantially above the median (142.0).",
        "The gap between mean and median underscores the pronounced operational asymmetry between low-demand overnight intervals and extreme daytime rush-hour surges."
    )
    
    # Figure 3
    insert_figure(
        doc,
        os.path.join(viz_dir, "hourly_demand.png"),
        3,
        "Average Hourly Bike Rental Demand: Working Days vs. Non-Working Days",
        "Compare diurnal hourly ridership profiles between working days and weekends/holidays.",
        "Working days display a distinct bimodal commuter curve peaking sharply at 08:00 (368.6 rentals/hr) and 17:00 (461.5 rentals/hr). Non-working days follow a unimodal bell curve peaking during the afternoon (12:00-16:00, ~370-380 rentals/hr).",
        "Ridership is predominantly commuter-driven on weekdays and leisure/tourist-driven on weekends, requiring distinct station rebalancing schedules."
    )
    
    # Figure 4
    insert_figure(
        doc,
        os.path.join(viz_dir, "monthly_demand.png"),
        4,
        "Average Hourly Bike Rental Demand by Month (2011 vs. 2012)",
        "Examine monthly seasonality and quantify year-over-year fleet adoption growth.",
        "Ridership exhibits consistent sinusoidal seasonality peaking in June-September and bottoming in January. In every single month, 2012 demand substantially exceeded 2011 demand, resulting in an overall +64.88% annual ridership surge.",
        "Growth reflects rapid network expansion, increased station density, and maturing consumer habit formation in Washington D.C."
    )
    
    # Figure 5
    insert_figure(
        doc,
        os.path.join(viz_dir, "seasonal_demand.png"),
        5,
        "Hourly Bike Rental Distribution Across Seasons",
        "Evaluate the distribution, median, IQR, and mean of hourly rentals across the four calendar seasons.",
        "Fall (mean: 236.0 rentals/hr) and Summer (mean: 208.4 rentals/hr) record the highest demand, while Winter records 198.9 rentals/hr and Spring records the lowest (111.1 rentals/hr).",
        "Spring in the UCI dataset encompasses January to March where cold winter temperatures suppress early-year cycling activity."
    )
    
    # Figure 6
    insert_figure(
        doc,
        os.path.join(viz_dir, "weather_demand.png"),
        6,
        "Impact of Ambient Temperature and Weather Severity on Rental Demand",
        "Investigate the bivariate relationship between temperature, weather condition severity, and demand.",
        "Demand shows a strong positive linear relationship with ambient temperature (r = +0.40). Clear/few clouds weather supports high demand across all temperatures, whereas rain/snow causes severe drops.",
        "Temperature acts as the primary positive environmental catalyst, whereas precipitation acts as an acute deterrent."
    )
    
    # Figure 7
    insert_figure(
        doc,
        os.path.join(viz_dir, "workingday_demand.png"),
        7,
        "User Segmentation Demand: Working Days vs. Non-Working Days",
        "Quantify the compositional split between Casual and Registered riders across day types.",
        "On working days, registered riders average 167.6 rentals/hr (86.7% of demand) while casual riders account for 25.6 rentals/hr. On weekends/holidays, casual demand more than doubles to 57.4 rentals/hr (31.7% of total).",
        "Registered commuters provide reliable baseline revenue, while casual users represent high-margin weekend expansion opportunities."
    )
    
    # Figure 8
    insert_figure(
        doc,
        os.path.join(viz_dir, "correlation_heatmap.png"),
        8,
        "Pearson Correlation Heatmap of Environmental & Demand Features",
        "Quantify linear associations between continuous environmental variables and user demand tiers.",
        "Temperature and feeling temperature exhibit strong positive correlation with total rentals (r = +0.404 and +0.401). Humidity exhibits negative correlation (r = -0.323). Registered rentals correlate overwhelmingly with total demand (r = +0.972).",
        "Thermal comfort and absence of precipitation are the strongest environmental predictors of fleet utilization."
    )

    doc.add_page_break()

    # -------------------------------------------------------------
    # SECTION 10: KEY INSIGHTS
    # -------------------------------------------------------------
    add_heading_1(doc, "10. KEY QUANTIFIED INSIGHTS")
    add_body_p(doc, "Based on rigorous statistical computation and multi-dimensional exploratory analysis, the following eight core evidence-based findings are established:")
    
    insights = [
        ("Total Fleet Volume & Registered Core", "The Capital Bikeshare system logged 3,292,679 total rentals across 2011–2012. Registered annual subscribers constitute 81.17% (2,672,662 rides) of total demand, while casual recreational users represent 18.83% (620,017 rides). Registered commuters form the operational financial backbone."),
        ("Year-Over-Year Network Expansion", "System demand expanded by 64.88% year-over-year, growing from 1,243,103 rentals in 2011 (mean: 143.76 rentals/hr) to 2,049,576 rentals in 2012 (mean: 234.67 rentals/hr). This confirms rapid network adoption and service maturation."),
        ("Diurnal Commuter Dual-Peak Dynamic", "On working days, hourly demand exhibits a pronounced bimodal distribution peaking at 08:00 (mean: 368.6 rentals/hr) and 17:00 (mean: 461.5 rentals/hr), driven by office commutes. The lowest demand occurs at 04:00 (mean: 6.4 rentals/hr)."),
        ("Divergent Weekend vs. Weekday Behaviors", "On working days, registered riders dominate (167.6 rentals/hr vs. 25.6 for casual). Conversely, on weekends and holidays, casual demand more than doubles to 57.4 rentals/hr with a smooth afternoon peak (12:00–16:00)."),
        ("Seasonal Climatological Trajectory", "Ridership peaks in Fall (mean: 236.0 rentals/hr, 1,061,129 total rides) and Summer (mean: 208.4 rentals/hr), whereas Spring records the annual trough (mean: 111.1 rentals/hr), primarily due to cold Q1 winter weather."),
        ("Monthly Demand Range", "Monthly demand peaks in June (240.5 rentals/hr) and September (240.8 rentals/hr), while January records the annual minimum (94.4 rentals/hr), representing a 2.55-fold seasonal expansion."),
        ("Temperature as Strongest Environmental Catalyst", "Ambient temperature is the single strongest environmental driver of ridership with a Pearson correlation of r = +0.404 (feeling temperature r = +0.401). Warmer temperatures consistently stimulate higher demand across both user tiers."),
        ("Precipitation & Humidity Deterrence", "Relative humidity correlates negatively with demand (r = -0.323). Average hourly demand drops by 62.4% from Clear conditions (204.9 rentals/hr) to Light Rain/Snow (111.6 rentals/hr), with severe storms seeing virtually zero rides (36.0 rentals/hr).")
    ]
    
    for idx, (title, desc) in enumerate(insights, 1):
        add_bullet_p(doc, f"{desc}", bold_prefix=f"{idx}. {title.upper()}: ")

    # -------------------------------------------------------------
    # SECTION 11, 12, 13, 14: LIMITATIONS, FUTURE SCOPE, CONCLUSION, REFS
    # -------------------------------------------------------------
    add_heading_1(doc, "11. LIMITATIONS OF THE STUDY")
    add_bullet_p(doc, "The dataset is confined to the Washington, D.C. Capital Bikeshare system; findings may not directly generalize to cities with different topographies, cycling infrastructure, or weather profiles.", bold_prefix="• Geographic Specificity: ")
    add_bullet_p(doc, "The observations span 2011–2012, capturing early-stage bikeshare adoption before the advent of dockless e-bikes and modern micro-mobility apps.", bold_prefix="• Temporal Window: ")
    add_bullet_p(doc, "Observed correlations between environmental variables and ridership do not prove causal relationships; confounding factors like daylight hours and public events contribute.", bold_prefix="• Descriptive Scope: ")
    add_bullet_p(doc, "The dataset aggregates system-wide hourly totals without individual station GPS coordinates, dock availability constraints, or trip origin-destination matrices.", bold_prefix="• Spatial Aggregation: ")

    add_heading_1(doc, "12. FUTURE SCOPE AND PREDICTIVE OPPORTUNITIES")
    add_bullet_p(doc, "Develop Random Forest, Gradient Boosting (XGBoost/LightGBM), and regularized Ridge regression models to predict hourly demand based on weather forecasts.", bold_prefix="• Supervised Regression: ")
    add_bullet_p(doc, "Implement SARIMAX, Prophet, and LSTM neural architectures to model multi-day ahead ridership patterns and capture holiday anomalies.", bold_prefix="• Time-Series Forecasting: ")
    add_bullet_p(doc, "Formulate mixed-integer linear programming (MILP) models for optimal van routing to prevent station dock starvation during morning and evening rush hours.", bold_prefix="• Dynamic Fleet Rebalancing: ")

    add_heading_1(doc, "13. CONCLUSION")
    add_body_p(doc, "This Week 1 project successfully executed the complete data preparation and exploratory analysis lifecycle for the UCI Bike Sharing Dataset. Starting from programmatic acquisition of raw data, we verified dataset completeness (0 missing values, 0 duplicate records), verified domain assertions, engineered intuitive temporal and physical features, and computed comprehensive descriptive statistics.")
    add_body_p(doc, "The exploratory analysis established that bike rental demand in Washington D.C. is governed by three primary forces: diurnal commute schedules (bimodal 08:00 and 17:00 peaks on working days), annual sinusoidal seasonality (summer/fall peaks vs winter troughs), and environmental comfort (positive temperature correlation r = +0.40, negative humidity/rain deterrence). These empirical findings provide the critical foundation for future machine learning and predictive operations.")

    add_heading_1(doc, "14. REFERENCES")
    add_bullet_p(doc, "Fanaee-T, H., & Gama, J. (2014). Event labeling combining ensemble detectors and background knowledge. Progress in Artificial Intelligence, 2(2), 113-127. DOI: 10.1007/s13748-013-0040-3.")
    add_bullet_p(doc, "UCI Machine Learning Repository. Bike Sharing Dataset. DOI: 10.24432/C5W894. URL: https://archive.ics.uci.edu/dataset/275/bike+sharing+dataset.")
    add_bullet_p(doc, "McKinney, W. (2010). Data Structures for Statistical Computing in Python. Proceedings of the 9th Python in Science Conference, 51-56.")
    add_bullet_p(doc, "Harris, C. R., et al. (2020). Array programming with NumPy. Nature, 585(7825), 357-362.")
    add_bullet_p(doc, "Hunter, J. D. (2007). Matplotlib: A 2D graphics environment. Computing in Science & Engineering, 9(3), 90-95.")
    add_bullet_p(doc, "Waskom, M. L. (2021). Seaborn: statistical data visualization. Journal of Open Source Software, 6(60), 3021.")
    add_bullet_p(doc, "Tukey, J. W. (1977). Exploratory Data Analysis. Addison-Wesley.")

    doc.add_page_break()

    # -------------------------------------------------------------
    # APPENDICES
    # -------------------------------------------------------------
    add_heading_1(doc, "APPENDIX A: COMPLETE PYTHON SOURCE ARCHITECTURE")
    add_body_p(doc, "The project is architected into modular, reusable Python scripts under the 'src/' package, complemented by an interactive, fully executed Jupyter Notebook in 'notebooks/Week1_Bike_Sharing_EDA.ipynb':")
    add_bullet_p(doc, "Handles automated downloading, archive extraction, and raw data integrity validation.", bold_prefix="• src/data_acquisition.py: ")
    add_bullet_p(doc, "Executes data quality audits, boundary constraint assertions, type conversion, and feature engineering.", bold_prefix="• src/data_cleaning.py: ")
    add_bullet_p(doc, "Calculates parametric/non-parametric statistics, aggregations, correlations, and insight generation.", bold_prefix="• src/exploratory_analysis.py: ")
    add_bullet_p(doc, "Generates all 8 publication-grade 300 DPI visualization artifacts.", bold_prefix="• src/visualization.py: ")
    add_bullet_p(doc, "Compiles and formats this comprehensive academic Word document report.", bold_prefix="• src/generate_docx_report.py: ")

    add_heading_1(doc, "APPENDIX B: DATA QUALITY AUDIT & TRANSFORMATION MATRIX")
    
    # Data Quality Report Table from CSV
    df_dq = pd.read_csv(os.path.join(base_dir, "outputs", "data_quality_report.csv"))
    tbl_dq = doc.add_table(rows=len(df_dq) + 1, cols=6)
    tbl_dq.alignment = WD_TABLE_ALIGNMENT.CENTER
    dq_headers = ["Feature", "Raw Type", "Cleaned Type", "Nulls", "Uniques", "Status"]
    
    for c_idx, h_text in enumerate(dq_headers):
        cell = tbl_dq.cell(0, c_idx)
        set_cell_background(cell, "1D3557")
        p = cell.paragraphs[0]
        format_paragraph(p, space_before=3, space_after=3)
        r = p.add_run(h_text)
        r.font.name = "Calibri"
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    for r_idx, row in df_dq.iterrows():
        vals = [
            str(row["Column_Name"]), str(row["Raw_Data_Type"]), str(row["Cleaned_Data_Type"]),
            str(row["Null_Count"]), str(row["Unique_Values"]), str(row["Validation_Status"])
        ]
        for c_idx, val in enumerate(vals):
            cell = tbl_dq.cell(r_idx + 1, c_idx)
            set_cell_background(cell, "FFFFFF" if r_idx % 2 == 0 else "F8F9FA")
            p = cell.paragraphs[0]
            format_paragraph(p, space_before=2, space_after=2)
            r = p.add_run(val)
            r.font.name = "Calibri"
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(40, 40, 40)
            
    set_table_borders(tbl_dq, color="D0D7DE")
    
    # Save Report
    doc.save(output_docx_path)
    file_size = os.path.getsize(output_docx_path)
    print(f"[SUCCESS] Academic DOCX report generated at: {output_docx_path}")
    print(f"[INFO] Report file size: {file_size:,} bytes ({file_size / (1024*1024):.2f} MB)")
    return output_docx_path

if __name__ == "__main__":
    build_docx_report()
